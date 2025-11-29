"""
Document Processing Service

Uses pdfplumber to extract text from PDF and other documents.
Supports PDF, Excel, CSV, DOCX documents, and plain text.
"""

import logging
import hashlib
from pathlib import Path
from typing import Optional, Tuple
from io import BytesIO

logger = logging.getLogger(__name__)

# Suppress non-critical pdfminer font warnings
logging.getLogger("pdfminer.pdffont").setLevel(logging.ERROR)


class DocumentProcessingService:
    """Service for processing documents with pdfplumber."""

    def __init__(self):
        """Initialize document processing service."""
        pass

    def process_document(
        self, file_content: bytes, filename: str, document_type: str
    ) -> Tuple[bool, Optional[str], Optional[str], Optional[str]]:
        """
        Process document and extract markdown content.

        Args:
            file_content: Binary content of the document
            filename: Original filename
            document_type: Type of document (pdf, excel, csv, word, txt)

        Returns:
            Tuple of (
                success: bool,
                markdown_content: str or None,
                content_hash: str or None,
                error_message: str or None
            )
        """
        try:
            # Calculate content hash
            content_hash = hashlib.sha256(file_content).hexdigest()

            # Handle different document types
            if document_type == "pdf":
                return self._process_pdf(file_content, content_hash)
            elif document_type == "txt":
                return self._process_txt(file_content, content_hash)
            elif document_type in ["excel", "csv"]:
                return self._process_spreadsheet(file_content, document_type, content_hash)
            elif document_type == "word":
                return self._process_word(file_content, content_hash)
            else:
                return False, None, None, f"Unsupported document type: {document_type}"

        except Exception as e:
            error_msg = f"Error processing document: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return False, None, None, error_msg

    def _process_pdf(
        self, file_content: bytes, content_hash: str
    ) -> Tuple[bool, Optional[str], Optional[str], Optional[str]]:
        """Process PDF document with pdfplumber."""
        try:
            import pdfplumber

            # Open PDF from bytes
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                # Extract text from all pages
                text_parts = []
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"## Page {page_num}\n\n{text}\n")

                if not text_parts:
                    fallback_markdown = (
                        "## Page 1\n\n"
                        "_В PDF не обнаружен текстовый слой. Документ загружен, но текст недоступен для извлечения._\n"
                    )
                    logger.warning("PDF has no extractable text; storing fallback notice as markdown content.")
                    markdown_content = fallback_markdown
                else:
                    # Combine all pages into markdown
                    markdown_content = "\n".join(text_parts)

                logger.info(
                    f"Successfully processed PDF, extracted {len(markdown_content)} characters from {len(pdf.pages)} pages"
                )
                return True, markdown_content, content_hash, None

        except Exception as e:
            error_msg = f"Error processing PDF: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return False, None, None, error_msg

    def _process_txt(
        self, file_content: bytes, content_hash: str
    ) -> Tuple[bool, Optional[str], Optional[str], Optional[str]]:
        """Process plain text file."""
        try:
            # Decode text
            text_content = file_content.decode("utf-8")

            # Convert to markdown (just wrap in code block or keep as is)
            markdown_content = f"```\n{text_content}\n```"

            logger.info(f"Successfully processed TXT file, {len(text_content)} characters")
            return True, markdown_content, content_hash, None

        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                try:
                    text_content = file_content.decode(encoding)
                    markdown_content = f"```\n{text_content}\n```"
                    return True, markdown_content, content_hash, None
                except:
                    continue

            return False, None, None, "Failed to decode text file"

    def _process_spreadsheet(
        self, file_content: bytes, doc_type: str, content_hash: str
    ) -> Tuple[bool, Optional[str], Optional[str], Optional[str]]:
        """Process Excel or CSV file."""
        try:
            import pandas as pd

            # Read file
            if doc_type == "csv":
                df = pd.read_csv(BytesIO(file_content))
            else:  # excel
                df = pd.read_excel(BytesIO(file_content))

            # Convert to markdown table
            markdown_content = df.to_markdown(index=False)

            logger.info(f"Successfully processed {doc_type.upper()} file, {len(df)} rows")
            return True, markdown_content, content_hash, None

        except Exception as e:
            error_msg = f"Error processing spreadsheet: {str(e)}"
            logger.error(error_msg)
            return False, None, None, error_msg

    def _process_word(
        self, file_content: bytes, content_hash: str
    ) -> Tuple[bool, Optional[str], Optional[str], Optional[str]]:
        """Process Word document (.docx)."""
        try:
            from docx import Document

            # Open Word document from bytes
            doc = Document(BytesIO(file_content))

            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            if not text_parts:
                return False, None, None, "No text content found in Word document"

            # Convert to markdown
            markdown_content = "\n\n".join(text_parts)

            logger.info(f"Successfully processed Word document, extracted {len(markdown_content)} characters")
            return True, markdown_content, content_hash, None

        except Exception as e:
            error_msg = f"Error processing Word document: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return False, None, None, error_msg

    @staticmethod
    def detect_document_type(filename: str) -> str:
        """
        Detect document type from filename.

        Args:
            filename: Original filename

        Returns:
            Document type: pdf, excel, csv, word, txt, or other
        """
        extension = Path(filename).suffix.lower()

        type_mapping = {
            ".pdf": "pdf",
            ".xlsx": "excel",
            ".xls": "excel",
            ".csv": "csv",
            ".docx": "word",
            # Legacy binary .doc files are unsupported – python-docx cannot parse them reliably
            ".txt": "txt",
        }

        return type_mapping.get(extension, "other")


# Singleton instance
document_processing_service = DocumentProcessingService()
